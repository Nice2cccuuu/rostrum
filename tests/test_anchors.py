"""Tests for click-to-select.

The hard part of this feature is not hit-testing arithmetic -- it is that the
geometry being tested is *computed*, not read back from the file. A bullet is a
paragraph inside a shared text frame and has no box of its own in the package, so
anchors are derived from the same font metrics the renderer used to lay the text
out. If that derivation drifts, hit-testing confidently returns the wrong uid.

So the tests here fall into two groups. Pure geometry (containment, ranking,
coverage) is checked against hand-built anchor maps, where the expected answer is
unarguable. The derived geometry is checked against a real render, by asserting
the properties that were actually violated during development: boxes ordered down
the page, sized in whole lines, and positioned according to the *resolved*
vertical anchor rather than a default.

That last one is the reason this file exists. Every box was 150 pixels above its
text because ``vertical_anchor`` returns ``None`` on a shape that inherits it, and
no assertion about anchor counts or box arithmetic would ever have caught it.
"""

from __future__ import annotations

import itertools

import pytest

from rostrum.ir.enums import Density, Scenario
from rostrum.render.anchors import (
    Anchor,
    AnchorMap,
    Box,
    anchors_to_json,
    hit_test,
    lasso,
    resolve_selection,
)

pytest.importorskip("pptx")
pytest.importorskip("docx")


# --------------------------------------------------------------------------- #
# Geometry
# --------------------------------------------------------------------------- #


def test_a_box_knows_what_it_contains():
    box = Box(x=0.1, y=0.2, w=0.3, h=0.4)
    assert box.contains(0.2, 0.3)
    assert box.contains(0.1, 0.2)  # boundaries count
    assert not box.contains(0.05, 0.3)
    assert not box.contains(0.2, 0.7)


def test_distance_is_zero_inside_and_grows_outside():
    box = Box(x=0.4, y=0.4, w=0.2, h=0.2)
    assert box.distance_to(0.5, 0.5) == 0.0
    assert box.distance_to(0.5, 0.35) == pytest.approx(0.05)
    assert box.distance_to(0.3, 0.3) == pytest.approx(0.1 * 2**0.5)


def test_coverage_is_measured_against_the_anchor_not_the_lasso():
    """Asymmetric on purpose: a loose circle around a bullet still selects it."""
    bullet = Box(x=0.1, y=0.5, w=0.8, h=0.05)
    generous = Box(x=0.0, y=0.4, w=1.0, h=0.3)
    assert bullet.overlap(generous) == pytest.approx(1.0)
    assert generous.overlap(bullet) < 0.2


def _map(*anchors: Anchor) -> AnchorMap:
    amap = AnchorMap(deck_uid="dck_x", slide_width_pt=720, slide_height_pt=405)
    for a in anchors:
        amap.add(a)
    return amap


def _anchor(uid, kind, x, y, w, h, *, slide=0, conf=1.0, preview="") -> Anchor:
    return Anchor(
        uid=uid,
        kind=kind,
        slide_index=slide,
        slide_uid="sld_1",
        path=uid,
        box=Box(x=x, y=y, w=w, h=h),
        preview=preview,
        confidence=conf,
    )


# --------------------------------------------------------------------------- #
# Hit testing
# --------------------------------------------------------------------------- #


def test_a_click_finds_the_bullet_not_the_box_around_it():
    """Otherwise every click resolves to the body placeholder."""
    amap = _map(
        _anchor("sld_1", "slide", 0, 0, 1, 1),
        _anchor("blk_1", "block", 0.1, 0.4, 0.8, 0.08),
    )
    result = hit_test(amap, 0, 0.5, 0.44)
    assert result.best.uid == "blk_1"


def test_a_near_miss_still_finds_the_nearest_target():
    """Clicking in the gap between bullets must not report nothing."""
    amap = _map(_anchor("blk_1", "block", 0.1, 0.4, 0.8, 0.05))
    result = hit_test(amap, 0, 0.5, 0.39, tolerance=0.02)
    assert result.best is not None
    assert result.hits[0].why.startswith("距离")


def test_a_click_far_from_everything_finds_nothing():
    amap = _map(_anchor("blk_1", "block", 0.1, 0.4, 0.8, 0.05))
    assert not hit_test(amap, 0, 0.5, 0.9, tolerance=0.02).hits


def test_a_click_between_two_bullets_is_reported_as_ambiguous():
    """A near-tie must be surfaced, not silently resolved.

    Picking one of two adjacent bullets without saying so is how a user ends up
    editing a sentence they did not point at.
    """
    amap = _map(
        _anchor("blk_1", "block", 0.1, 0.40, 0.8, 0.05),
        _anchor("blk_2", "block", 0.1, 0.45, 0.8, 0.05),
    )
    assert hit_test(amap, 0, 0.5, 0.4499).ambiguous


def test_a_confident_hit_is_not_ambiguous():
    amap = _map(
        _anchor("sld_1", "slide", 0, 0, 1, 1),
        _anchor("blk_1", "block", 0.1, 0.40, 0.8, 0.05),
    )
    assert not hit_test(amap, 0, 0.5, 0.42).ambiguous


def test_a_click_on_empty_space_falls_back_to_the_page():
    amap = _map(
        _anchor("sld_1", "slide", 0, 0, 1, 1),
        _anchor("blk_1", "block", 0.1, 0.1, 0.8, 0.05),
    )
    result = hit_test(amap, 0, 0.5, 0.8)
    assert result.best.kind == "slide"


def test_clicks_are_scoped_to_one_page():
    amap = _map(
        _anchor("blk_1", "block", 0.1, 0.4, 0.8, 0.05, slide=0),
        _anchor("blk_2", "block", 0.1, 0.4, 0.8, 0.05, slide=1),
    )
    assert hit_test(amap, 1, 0.5, 0.42).best.uid == "blk_2"


def test_lower_confidence_geometry_ranks_lower():
    """Computed boxes should lose to measured ones when both match."""
    amap = _map(
        _anchor("exact", "block", 0.1, 0.4, 0.8, 0.05, conf=1.0),
        _anchor("derived", "block", 0.1, 0.4, 0.8, 0.05, conf=0.5),
    )
    assert hit_test(amap, 0, 0.5, 0.42).best.uid == "exact"


# --------------------------------------------------------------------------- #
# Lasso
# --------------------------------------------------------------------------- #


def test_a_lasso_selects_everything_it_covers():
    amap = _map(
        _anchor("blk_1", "block", 0.1, 0.40, 0.8, 0.05),
        _anchor("blk_2", "block", 0.1, 0.46, 0.8, 0.05),
        _anchor("blk_3", "block", 0.1, 0.80, 0.8, 0.05),
    )
    result = lasso(amap, 0, Box(x=0.05, y=0.35, w=0.9, h=0.2))
    assert result.uids == ["blk_1", "blk_2"]


def test_a_lasso_ignores_the_page_anchor():
    """A rectangle means "these things", never "this page"."""
    amap = _map(
        _anchor("sld_1", "slide", 0, 0, 1, 1),
        _anchor("blk_1", "block", 0.1, 0.4, 0.8, 0.05),
    )
    assert lasso(amap, 0, Box(x=0.0, y=0.0, w=1.0, h=1.0)).uids == ["blk_1"]


def test_a_lasso_that_merely_clips_a_line_does_not_select_it():
    amap = _map(_anchor("blk_1", "block", 0.1, 0.40, 0.8, 0.10))
    assert not lasso(amap, 0, Box(x=0.1, y=0.48, w=0.8, h=0.03)).hits


def test_a_lasso_returns_results_in_reading_order():
    """A selection is a list the user will see; top-to-bottom is how they drew it."""
    amap = _map(
        _anchor("lower", "block", 0.1, 0.60, 0.8, 0.05),
        _anchor("upper", "block", 0.1, 0.40, 0.8, 0.05),
    )
    result = lasso(amap, 0, Box(x=0.05, y=0.35, w=0.9, h=0.35))
    assert result.uids == ["upper", "lower"]


def test_a_lasso_is_never_ambiguous():
    """Every match was intended; warning about equal scores trains users to ignore."""
    amap = _map(
        _anchor("blk_1", "block", 0.1, 0.40, 0.8, 0.05),
        _anchor("blk_2", "block", 0.1, 0.46, 0.8, 0.05),
    )
    result = lasso(amap, 0, Box(x=0.05, y=0.35, w=0.9, h=0.2))
    assert len(result.hits) == 2
    assert not result.ambiguous


def test_resolve_selection_handles_both_gestures():
    amap = _map(_anchor("blk_1", "block", 0.1, 0.4, 0.8, 0.05))
    assert resolve_selection(amap, slide=0, point=(0.5, 0.42)).best.uid == "blk_1"
    assert resolve_selection(
        amap, slide=0, rect=(0.05, 0.35, 0.9, 0.2)
    ).uids == ["blk_1"]
    with pytest.raises(ValueError, match="point or a rect"):
        resolve_selection(amap, slide=0)


# --------------------------------------------------------------------------- #
# Derived geometry, against a real render
# --------------------------------------------------------------------------- #


@pytest.fixture(scope="module")
def rendered(tmp_path_factory):
    import docx

    from rostrum.budget.allocate import allocate
    from rostrum.ingest.docx_parser import parse_docx
    from rostrum.ingest.planner import plan_deck
    from rostrum.render import render_pptx
    from rostrum.templates import bind, capacity_caps, ingest_pptx
    from rostrum.themes import DEFAULT_THEME_ID, build_template, get_theme

    d = tmp_path_factory.mktemp("anch")
    doc = docx.Document()
    doc.core_properties.title = "锚点测试"
    doc.add_heading("锚点测试", 1)
    doc.add_heading("研究背景与问题", 1)
    doc.add_paragraph(
        "深度表征学习在大规模标注数据上取得了显著成功，但在医学影像、工业质检等"
        "真实场景中，获取高质量标注的成本极高，可用样本常低于千级规模。"
    )
    doc.add_paragraph("现有自监督方法在小样本条件下普遍出现表征坍缩。")
    # A deliberately long single bullet, pinned so planning keeps it intact. The
    # anchor geometry test needs at least one bullet that wraps onto a second
    # line; without one, every box is the same height and the test cannot tell
    # correct wrapping from ignored wrapping.
    doc.add_paragraph(
        "该现象在跨域迁移与长尾分布并存的条件下尤为显著且难以通过简单的数据增强"
        "或损失函数调整加以缓解因此需要从表征空间的几何结构入手重新设计约束项"
    )
    doc.add_paragraph("本项目拟解决的核心科学问题是构造不坍缩的表征空间。")
    doc.add_heading("创新点", 1)
    doc.add_paragraph("提出结构一致性正则", style="List Bullet")
    doc.add_paragraph("参数量降低54%", style="List Bullet")
    doc.add_heading("研究基础", 1)
    doc.add_paragraph("前期工作已在三个数据集上验证。")
    doc.add_heading("可行性分析", 1)
    doc.add_paragraph("团队具备算力条件。")
    src = d / "m.docx"
    doc.save(str(src))

    template = d / "t.pptx"
    build_template(get_theme(DEFAULT_THEME_ID), str(template))

    parsed = parse_docx(str(src), asset_dir=str(d / "a"))
    # Compact density keeps long bullets long (42 units per bullet rather than
    # 18), which this fixture needs: anchor geometry is only interesting when some
    # bullets wrap onto a second line, and sparser settings now split prose into
    # short single-line points.
    deck = plan_deck(
        parsed,
        total_seconds=480,
        scenario=Scenario.GRANT_DEFENSE,
        density=Density.COMPACT,
    )
    contract, _ = ingest_pptx(str(template), template_id="t", license="builtin")
    binding = bind(deck, contract)
    allocate(deck, apply=True, capacity=capacity_caps(binding))
    report = render_pptx(deck, contract, binding, str(d / "out.pptx"))
    return deck, report


def test_rendering_produces_an_anchor_for_every_visible_block(rendered):
    deck, report = rendered
    amap = report.anchors
    assert amap is not None

    from rostrum.ir.enums import Channel

    on_slide = {
        b.uid
        for _, _, b in deck.iter_blocks()
        if b.channel is Channel.SLIDE and b.content.strip()
    }
    anchored = {a.uid for a in amap.anchors if a.kind == "block"}
    missing = on_slide - anchored - set(report.demoted_to_notes)
    assert not missing, f"{len(missing)} visible block(s) cannot be clicked"


def test_every_page_has_a_page_level_anchor(rendered):
    _, report = rendered
    for index in range(report.anchors.slide_count):
        kinds = {a.kind for a in report.anchors.for_slide(index)}
        assert "slide" in kinds


def test_block_anchors_run_down_the_page_in_order(rendered):
    """Boxes must follow the reading order of the blocks that produced them."""
    _, report = rendered
    for index in range(report.anchors.slide_count):
        blocks = [a for a in report.anchors.for_slide(index) if a.kind == "block"]
        tops = [a.box.y for a in blocks]
        assert tops == sorted(tops), f"page {index + 1} anchors are out of order"


def test_block_anchors_do_not_overlap_each_other(rendered):
    """Overlapping boxes make a click between two bullets unresolvable."""
    _, report = rendered
    for index in range(report.anchors.slide_count):
        blocks = sorted(
            (a for a in report.anchors.for_slide(index) if a.kind == "block"),
            key=lambda a: a.box.y,
        )
        for first, second in itertools.pairwise(blocks):
            assert first.box.y2 <= second.box.y + 1e-6


def test_block_anchors_sit_below_the_title(rendered):
    """A body anchor overlapping the heading would steal clicks meant for it."""
    _, report = rendered
    for index in range(report.anchors.slide_count):
        page = report.anchors.for_slide(index)
        titles = [a for a in page if a.kind == "title"]
        blocks = [a for a in page if a.kind == "block"]
        if not titles or not blocks:
            continue
        assert min(b.box.y for b in blocks) >= min(t.box.y for t in titles)


def test_anchors_respect_the_inherited_vertical_alignment(rendered):
    """The defect this whole file was written for.

    The built-in themes centre body text vertically, and that setting lives on the
    layout, not the shape. Reading it from the shape returned ``None``, which was
    treated as top alignment and put every box 150px above its text -- correct
    arithmetic on the wrong input.

    With centring honoured, body anchors start well below the placeholder's top.
    """
    _, report = rendered
    pages = [
        i
        for i in range(report.anchors.slide_count)
        if len([a for a in report.anchors.for_slide(i) if a.kind == "block"]) >= 2
    ]
    assert pages, "fixture produced no multi-bullet page"

    for index in pages:
        blocks = [a for a in report.anchors.for_slide(index) if a.kind == "block"]
        first = min(a.box.y for a in blocks)
        last = max(a.box.y2 for a in blocks)
        text_middle = (first + last) / 2
        # Centred text straddles the middle of the page's body area rather than
        # hugging its top.
        assert 0.3 < text_middle < 0.8, (
            f"page {index + 1}: text centre at {text_middle:.3f} suggests the "
            "vertical anchor was not resolved"
        )


def test_a_multi_line_bullet_gets_a_taller_box(rendered):
    """Box height must reflect wrapping, or long bullets are half-clickable."""
    _, report = rendered
    blocks = [a for a in report.anchors.anchors if a.kind == "block"]
    heights = {round(a.box.h, 4) for a in blocks}
    assert len(heights) > 1, "every bullet got the same height; wrapping was ignored"


def test_derived_anchors_admit_they_are_derived(rendered):
    """Computed geometry must not claim the certainty of measured geometry."""
    _, report = rendered
    for anchor in report.anchors.anchors:
        if anchor.kind == "block":
            assert anchor.confidence < 1.0
        if anchor.kind in ("title", "slide"):
            assert anchor.confidence == 1.0


def test_anchors_survive_a_save_and_load(rendered, tmp_path):
    _, report = rendered
    path = tmp_path / "anchors.json"
    report.anchors.save(str(path))
    reloaded = AnchorMap.load(str(path))
    assert len(reloaded.anchors) == len(report.anchors.anchors)
    assert reloaded.deck_uid == report.anchors.deck_uid


def test_the_compact_json_form_is_usable_by_a_ui(rendered):
    import json

    _, report = rendered
    data = json.loads(anchors_to_json(report.anchors))
    assert data["anchors"]
    first = data["anchors"][0]
    assert len(first["box"]) == 4
    assert {"uid", "kind", "slide", "path", "box"} <= set(first)


def test_titles_and_pages_stay_distinguishable_in_a_candidate_list(rendered):
    """Both carry the slide's uid; a user must still be able to tell them apart."""
    _, report = rendered
    page = report.anchors.for_slide(2)
    paths = [a.path for a in page if a.kind in ("title", "slide")]
    assert len(paths) == len(set(paths))


# --------------------------------------------------------------------------- #
# Pointing feeds the same editing mechanism as language
# --------------------------------------------------------------------------- #


def test_a_click_resolves_to_a_uid_that_can_be_edited(rendered):
    """The whole point: pointing replaces describing, not the operation itself."""
    from rostrum.patch.session import Session

    deck, report = rendered
    target = next(
        a
        for a in report.anchors.anchors
        if a.kind == "block" and len(a.preview) > 10
    )
    hit = hit_test(report.anchors, target.slide_index, target.box.cx, target.box.cy)
    assert hit.best.uid == target.uid

    session = Session(original=deck)
    result, _diff, applied = session.say("这条改短一点", selection=[hit.best.uid])
    assert result.ok
    assert applied is not None
    assert result.patch.operations[0].target == target.uid


def test_a_lasso_selection_edits_every_item_it_covered(rendered):
    from rostrum.ir.enums import Channel
    from rostrum.patch.session import Session

    deck, report = rendered

    # An agenda page pins its items, and the interpreter rightly refuses to move
    # a pinned block without confirmation. Pick a page whose content is free.
    def unpinned(index: int) -> list:
        found = [a for a in report.anchors.for_slide(index) if a.kind == "block"]
        return [a for a in found if not deck.find(a.uid).pinned]

    page = next(
        i for i in range(report.anchors.slide_count) if len(unpinned(i)) >= 2
    )
    blocks = unpinned(page)
    top = min(a.box.y for a in blocks)
    bottom = max(a.box.y2 for a in blocks)
    selected = lasso(
        report.anchors, page, Box(x=0.0, y=top - 0.01, w=1.0, h=bottom - top + 0.02)
    )
    assert len(selected.uids) >= 2

    session = Session(original=deck)
    result, _, applied = session.say("这些放到讲稿里", selection=selected.uids)
    assert result.ok and applied is not None

    # Assert against the ops the patch actually emitted rather than against the
    # raw selection: figures cannot be spoken, so the interpreter legitimately
    # drops them from a move-to-script request and says so.
    moved = [o.target for o in result.patch.operations if o.op == "set_channel"]
    assert len(moved) >= 2
    for uid in moved:
        assert session.current.find(uid).channel is Channel.SCRIPT


def test_pointing_beats_describing_on_confidence(rendered):
    """A selection is what the user indicated; it should not be second-guessed."""
    from rostrum.patch.interpret import interpret

    deck, report = rendered
    target = next(a for a in report.anchors.anchors if a.kind == "block")

    pointed = interpret("这条改短一点", deck, selection=[target.uid])
    described = interpret("把第3页的字精简一下", deck)
    assert pointed.confidence >= described.confidence


def test_an_overlay_can_be_drawn_for_verification(rendered, tmp_path):
    """Computed geometry has to be checkable against the pixels it describes."""
    from PIL import Image

    from rostrum.render.anchors import draw_overlay

    _, report = rendered
    page = tmp_path / "page.png"
    Image.new("RGB", (1200, 675), "white").save(page)

    out = draw_overlay(report.anchors, 2, str(page), str(tmp_path / "o.png"))
    assert Image.open(out).size == (1200, 675)
