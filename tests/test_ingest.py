"""Tests for the ingest layer: parsers, offsets, and the content planner.

The load-bearing property under test is that **every span points at the text it
claims to**. A deck whose offsets have drifted looks correct until a user clicks
a bullet to edit it and lands somewhere else, so ``verify()`` is asserted after
every parse rather than only in a dedicated test.
"""

from __future__ import annotations

import importlib.util
import os

import pytest

from rostrum.ingest.docx_parser import parse_docx
from rostrum.ingest.latex_parser import _to_text, parse_latex
from rostrum.ingest.model import (
    ParsedDocument,
    Segment,
    SegmentKind,
    TextBuilder,
    normalize,
)
from rostrum.ingest.pdf_parser import _collapse_spaced_run, parse_pdf
from rostrum.ingest.planner import (
    _affiliation_of,
    _rubric_for,
    _score,
    _split_claim,
    plan_deck,
)
from rostrum.ir.enums import BlockType, Channel, Density, Scenario, SlideRole
from rostrum.ir.validate import Severity, validate

pytest.importorskip("docx")


# --------------------------------------------------------------------------- #
# Normalisation and the text/offset contract
# --------------------------------------------------------------------------- #


def test_normalize_strips_invisibles_but_keeps_visible_characters():
    raw = "研\u200b究\ufeff目标\u00a0与内容"
    out = normalize(raw)
    assert out == "研究目标 与内容"
    # Visible characters must survive untouched: folding smart quotes or
    # full-width forms here would make Derivation.VERBATIM a lie.
    assert normalize("“引号”和（全角）") == "“引号”和（全角）"


def test_normalize_collapses_whitespace_runs():
    assert normalize("a   b\t\tc") == "a b c"
    assert normalize("line\n\n\npara") == "line\n\npara"


def test_text_builder_offsets_address_the_normalized_text():
    b = TextBuilder()
    b.add("第一段内容", SegmentKind.PARAGRAPH)
    b.add("第二段内容", SegmentKind.PARAGRAPH)
    for seg in b.segments:
        assert b.text[seg.start : seg.end] == seg.text


def test_verify_detects_offset_drift():
    doc = ParsedDocument(
        doc_id="d",
        text="正确的文本内容",
        segments=[Segment(text="错误", kind=SegmentKind.PARAGRAPH, start=0, end=2)],
    )
    problems = doc.verify()
    assert problems, "verify must catch a span that does not match the text"


def test_sha256_tracks_normalized_text_not_the_file():
    a = ParsedDocument(doc_id="d", text="同样的内容")
    b = ParsedDocument(doc_id="d", text="同样的内容")
    c = ParsedDocument(doc_id="d", text="不同的内容")
    assert a.sha256 == b.sha256
    assert a.sha256 != c.sha256


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #


@pytest.fixture(scope="module")
def docx_path(tmp_path_factory):
    """A manuscript exercising every structure the parser must recognise."""
    import docx
    from docx.shared import Inches
    from PIL import Image

    d = tmp_path_factory.mktemp("docx")
    img = d / "fig.png"
    Image.new("RGB", (800, 500), "white").save(img)

    doc = docx.Document()
    doc.core_properties.title = "测试用申报书"
    doc.core_properties.author = "张三; 李四"

    doc.add_heading("测试用申报书", 1)
    doc.add_paragraph("张三，李四　某大学计算机学院")

    doc.add_heading("研究背景与问题", 1)
    doc.add_paragraph("现有方法在小样本条件下性能显著下降，这是本项目要解决的核心问题。")
    doc.add_paragraph("例如在医学影像领域，标注成本极高。此外数据分布也存在偏移。")

    doc.add_heading("创新点", 1)
    doc.add_paragraph("提出结构一致性正则项", style="List Bullet")
    doc.add_paragraph("精度提升7.5个百分点，参数量降低54%", style="List Bullet")

    doc.add_heading("研究目标与内容", 1)
    doc.add_paragraph("总体框架如下图所示。")
    doc.add_picture(str(img), width=Inches(4.0))
    doc.add_paragraph("图1 整体技术框架")

    doc.add_heading("研究基础", 1)
    doc.add_paragraph("表1 对比结果")
    t = doc.add_table(rows=3, cols=3)
    for j, v in enumerate(["方法", "准确率", "参数量"]):
        t.cell(0, j).text = v
    for j, v in enumerate(["Baseline", "0.812", "24M"]):
        t.cell(1, j).text = v
    for j, v in enumerate(["本文", "0.887", "11M"]):
        t.cell(2, j).text = v

    doc.add_heading("参考文献", 1)
    doc.add_paragraph("[1] Chen T, et al. ICML 2020.")

    path = d / "proposal.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture(scope="module")
def parsed_docx(docx_path, tmp_path_factory):
    out = tmp_path_factory.mktemp("assets")
    return parse_docx(docx_path, asset_dir=str(out))


def test_docx_offsets_are_exact(parsed_docx):
    assert parsed_docx.verify() == []


def test_docx_reads_metadata(parsed_docx):
    assert parsed_docx.title == "测试用申报书"
    assert parsed_docx.authors == ["张三", "李四"]


def test_docx_recovers_heading_hierarchy(parsed_docx):
    headings = [s.text for s in parsed_docx.segments if s.kind is SegmentKind.HEADING]
    assert "研究背景与问题" in headings
    assert "创新点" in headings


def test_docx_assigns_section_paths(parsed_docx):
    inside = [
        s
        for s in parsed_docx.segments
        if s.kind is SegmentKind.LIST_ITEM and "创新点" in s.section_path
    ]
    assert inside, "list items must record the section they belong to"


def test_docx_isolates_references(parsed_docx):
    refs = [s for s in parsed_docx.segments if s.kind is SegmentKind.REFERENCE]
    assert len(refs) == 1
    # A bibliography must never be mistaken for body prose; it would otherwise
    # compete for slide space against the author's actual claims.
    assert "Chen" in refs[0].text


def test_docx_extracts_images_as_original_bytes(parsed_docx):
    figures = [a for a in parsed_docx.assets if a.kind.value == "figure"]
    assert len(figures) == 1
    assert figures[0].path and os.path.exists(figures[0].path)
    assert figures[0].intrinsic_aspect == pytest.approx(800 / 500, rel=0.01)


def test_docx_binds_captions_in_both_directions(parsed_docx):
    """Figure captions follow their figure; table captions precede their table."""
    by_id = {a.asset_id: a for a in parsed_docx.assets}
    figure = next(a for a in by_id.values() if a.kind.value == "figure")
    table = next(a for a in by_id.values() if a.kind.value == "table")
    assert figure.caption and "整体技术框架" in figure.caption
    assert table.caption and "对比结果" in table.caption


def test_docx_captures_table_values_not_a_picture(parsed_docx):
    table = next(a for a in parsed_docx.assets if a.kind.value == "table")
    assert table.data is not None
    assert table.data["columns"] == ["方法", "准确率", "参数量"]
    assert ["Baseline", "0.812", "24M"] in table.data["rows"]


# --------------------------------------------------------------------------- #
# LaTeX
# --------------------------------------------------------------------------- #


@pytest.fixture(scope="module")
def tex_path(tmp_path_factory):
    d = tmp_path_factory.mktemp("tex")
    (d / "figures").mkdir()
    from PIL import Image

    Image.new("RGB", (640, 400), "white").save(d / "figures" / "pipeline.png")

    (d / "paper.tex").write_text(
        r"""
\documentclass{article}
\title{Representation Learning}
\author{Zhang San \and Li Si}
\begin{document}
\maketitle
% this comment must vanish
\begin{abstract}
Skipped.
\end{abstract}
\section{Introduction}
Existing methods collapse when $n$ is small, degrading $\mathcal{R}(h)$.
\section{Method}
\label{sec:method}
We constrain the geometry.
\begin{figure}
\includegraphics[width=0.8\textwidth]{figures/pipeline}
\caption{The proposed pipeline.}
\label{fig:pipe}
\end{figure}
\begin{equation}
\label{eq:bound}
\mathcal{R}(h) \le \hat{\mathcal{R}}(h)
\end{equation}
\subsection{Contributions}
\begin{itemize}
  \item A regulariser.
  \begin{itemize}
    \item Verified on benchmarks.
  \end{itemize}
  \item A 54\% reduction.
\end{itemize}
\begin{table}
\caption{Results.}
\begin{tabular}{lr}
\toprule
Method & Accuracy \\
\midrule
Ours & 0.887 \\
\bottomrule
\end{tabular}
\end{table}
Confirmed~\cite{chen2020}.
\end{document}
""",
        encoding="utf-8",
    )
    return str(d / "paper.tex")


@pytest.fixture(scope="module")
def parsed_tex(tex_path):
    return parse_latex(tex_path)


def test_latex_offsets_are_exact(parsed_tex):
    assert parsed_tex.verify() == []


def test_latex_reads_metadata(parsed_tex):
    assert parsed_tex.title == "Representation Learning"
    assert parsed_tex.authors == ["Zhang San", "Li Si"]


def test_latex_preserves_inline_maths_verbatim(parsed_tex):
    intro = next(
        s for s in parsed_tex.segments if s.kind is SegmentKind.PARAGRAPH
    )
    # Flattening maths would destroy exactly the content a theory talk exists to
    # present, so the delimiters and commands must survive intact.
    assert "$\\mathcal{R}(h)$" in intro.text or "$n$" in intro.text


def test_latex_keeps_display_equations_but_strips_labels(parsed_tex):
    eq = next(s for s in parsed_tex.segments if s.kind is SegmentKind.EQUATION)
    assert "\\label" not in eq.text
    assert eq.locator == "eq:bound"
    assert "\\mathcal{R}(h)" in eq.text


def test_latex_records_section_labels_as_locators(parsed_tex):
    method = next(
        s
        for s in parsed_tex.segments
        if s.kind is SegmentKind.HEADING and s.text == "Method"
    )
    assert method.locator == "sec:method"


def test_latex_flattens_nested_lists_with_levels(parsed_tex):
    items = [s for s in parsed_tex.segments if s.kind is SegmentKind.LIST_ITEM]
    assert {s.level for s in items} == {0, 1}


def test_latex_resolves_graphics_without_extension(parsed_tex):
    figure = next(a for a in parsed_tex.assets if a.kind.value == "figure")
    assert figure.path and figure.path.endswith("pipeline.png")


def test_latex_recovers_tabular_values(parsed_tex):
    table = next(a for a in parsed_tex.assets if a.kind.value == "table")
    assert table.data["columns"] == ["Method", "Accuracy"]
    assert ["Ours", "0.887"] in table.data["rows"]


def test_latex_skips_abstract_and_comments(parsed_tex):
    joined = parsed_tex.text
    assert "Skipped" not in joined
    assert "this comment" not in joined


def test_to_text_drops_citations_but_keeps_prose():
    assert _to_text(r"Confirmed by prior work~\cite{smith}.") == "Confirmed by prior work ."
    assert _to_text(r"\textbf{bold} and \emph{italic}") == "bold and italic"


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #


@pytest.mark.parametrize(
    "raw,expect_joined",
    [
        ("R e p r e s e n t a t i o n", True),
        ("I C M L 2 0 2 0", True),
        ("we can do it on a set of ten", False),
        ("the proposed method for a set of tasks", False),
    ],
)
def test_pdf_spacing_repair_leaves_real_prose_alone(raw, expect_joined):
    """Repairing exploded glyphs must not weld genuine words together.

    A parser that over-collapses turns "we can do it" into "wecandoit", which is
    worse than the original defect because it is unrecoverable.
    """
    out = _collapse_spaced_run(raw)
    if expect_joined:
        assert " " not in out.strip()
    else:
        assert out == raw


#: PyMuPDF is an optional extra, so its absence must skip one test rather than the
#: file. Calling ``pytest.importorskip`` inside a ``skipif`` argument does the
#: latter: decorator arguments are evaluated while the module loads, and the
#: Skipped exception it raises aborts collection for everything below it. That cost
#: 57 unrelated tests -- reported as a single "1 skipped", so the loss was
#: invisible. It never showed up in development because this machine had PyMuPDF
#: installed; it took a clone into a clean environment to surface.
_HAS_FITZ = importlib.util.find_spec("fitz") is not None


@pytest.mark.skipif(not _HAS_FITZ, reason="PyMuPDF (fitz) not installed")
def test_pdf_parses_and_offsets_are_exact(tmp_path):
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((60, 80), "Representation Learning", fontsize=18)
    page.insert_text((60, 130), "Existing methods collapse in this regime.", fontsize=10)
    path = tmp_path / "p.pdf"
    doc.save(str(path))
    doc.close()

    parsed = parse_pdf(str(path), asset_dir=str(tmp_path / "a"))
    assert parsed.verify() == []
    assert parsed.segments
    # The largest text on page one is the title.
    assert parsed.title and "Representation" in parsed.title


# --------------------------------------------------------------------------- #
# Planner
# --------------------------------------------------------------------------- #


@pytest.mark.parametrize(
    "title,expected",
    [
        ("研究背景与问题", "motivation"),
        ("创新点", "innovation"),
        ("研究基础", "prior_work"),  # must not be captured by "objectives"
        ("研究目标与内容", "objectives"),
        ("技术路线", "methods"),
        ("经费预算", "budget"),
        ("Introduction", "motivation"),
        ("Related Work", "prior_work"),
        ("完全无关的标题", None),
    ],
)
def test_rubric_mapping(title, expected):
    assert _rubric_for(title) == expected


def test_score_prefers_claims_over_elaboration():
    claim = Segment(
        text="本项目首次提出结构一致性正则，精度提升7.5个百分点",
        kind=SegmentKind.PARAGRAPH,
        start=0,
        end=1,
    )
    aside = Segment(
        text="例如在医学影像领域，此外数据分布通常也存在偏移",
        kind=SegmentKind.PARAGRAPH,
        start=0,
        end=1,
    )
    assert _score(claim, "innovation") > _score(aside, "innovation")


def test_score_stays_within_the_ir_bound():
    """Importance must remain in [0, 1] no matter how many signals fire."""
    loud = Segment(
        text="本项目首次提出核心关键创新，显著突破，精度提升99%，达到state-of-the-art",
        kind=SegmentKind.LIST_ITEM,
        start=0,
        end=1,
    )
    assert 0.0 <= _score(loud, "innovation") <= 1.0


def test_split_claim_cuts_at_sentence_boundaries():
    text = "现有方法在小样本条件下性能显著下降。这主要是因为表征坍缩，具体而言判别性不足。"
    head, note = _split_claim(text)
    assert head.endswith("。")
    assert note and "表征坍缩" in note
    # Nothing may be lost: the remainder becomes what the presenter says.
    assert head + note == text


def test_split_claim_leaves_short_text_whole():
    head, note = _split_claim("提出结构一致性正则")
    assert head == "提出结构一致性正则"
    assert note is None


def test_affiliation_keeps_the_institution_name():
    assert _affiliation_of("张三，李四　某大学计算机学院") == "某大学计算机学院"
    assert _affiliation_of("Zhang, Li  Tsinghua University") == "Tsinghua University"


def test_planner_builds_a_valid_deck(parsed_docx):
    deck = plan_deck(parsed_docx, total_seconds=480, scenario=Scenario.GRANT_DEFENSE)
    report = validate(deck)
    assert [f for f in report.findings if f.severity is Severity.ERROR] == []
    assert deck.sections


def test_planner_excludes_front_matter_from_content(parsed_docx):
    """The byline is cover metadata, not a slide."""
    deck = plan_deck(parsed_docx, total_seconds=480)
    texts = [
        b.content
        for _, slide in deck.iter_slides()
        for b in slide.blocks
    ]
    assert not any("某大学计算机学院" in t for t in texts)
    assert deck.meta.affiliation == "某大学计算机学院"


def test_planner_keeps_a_float_and_its_caption_on_one_slide(parsed_docx):
    deck = plan_deck(parsed_docx, total_seconds=480)
    visual_slides = [
        slide
        for _, slide in deck.iter_slides()
        if any(b.type in (BlockType.FIGURE, BlockType.TABLE) for b in slide.blocks)
    ]
    for slide in visual_slides:
        visuals = [
            b for b in slide.blocks if b.type in (BlockType.FIGURE, BlockType.TABLE)
        ]
        # One float per slide, and it carries its caption -- not two slides, one
        # holding an anchor and one holding a caption.
        assert len(visuals) == 1
        assert visuals[0].asset_ref is not None


def test_planner_never_puts_placeholder_text_on_a_slide(parsed_docx):
    deck = plan_deck(parsed_docx, total_seconds=480)
    for _, slide in deck.iter_slides():
        for block in slide.blocks:
            assert not block.content.startswith("[Figure")
            assert not block.content.startswith("[表格")


def test_planner_weights_innovation_above_administration(parsed_docx):
    deck = plan_deck(parsed_docx, total_seconds=480, scenario=Scenario.GRANT_DEFENSE)
    weights = {s.rubric_key: s.weight for s in deck.sections}
    if "innovation" in weights and "prior_work" in weights:
        assert weights["innovation"] > weights["prior_work"]


def test_planner_orders_sections_for_a_defence(parsed_docx):
    """Manuscript order is not talk order."""
    deck = plan_deck(parsed_docx, total_seconds=480, scenario=Scenario.GRANT_DEFENSE)
    keys = [s.rubric_key for s in deck.sections if s.rubric_key]
    if "motivation" in keys and "prior_work" in keys:
        assert keys.index("motivation") < keys.index("prior_work")


def test_planner_routes_elaboration_to_the_script(parsed_docx):
    """Detail must survive as speech rather than being dropped."""
    deck = plan_deck(parsed_docx, total_seconds=480)
    all_blocks = [b for _, slide in deck.iter_slides() for b in slide.blocks]
    scripted = [b for b in all_blocks if b.channel is Channel.SCRIPT]
    with_notes = [b for b in all_blocks if b.speaker_note]
    assert scripted or with_notes, "low-importance prose must go somewhere, not vanish"


def test_planner_marks_every_block_with_provenance(parsed_docx):
    deck = plan_deck(parsed_docx, total_seconds=480)
    for _, slide in deck.iter_slides():
        for block in slide.blocks:
            assert block.spans, f"{block.uid} has no source span"
            for span in block.spans:
                assert span.doc_id == parsed_docx.doc_id


def test_planner_respects_density_for_bullet_count(parsed_docx):
    sparse = plan_deck(parsed_docx, total_seconds=480, density=Density.SPARSE)
    compact = plan_deck(parsed_docx, total_seconds=480, density=Density.COMPACT)

    def max_bullets(deck):
        return max(
            (
                sum(1 for b in slide.blocks if b.type is BlockType.BULLET)
                for _, slide in deck.iter_slides()
            ),
            default=0,
        )

    assert max_bullets(sparse) <= max_bullets(compact)


def test_planner_assigns_visual_roles(parsed_docx):
    deck = plan_deck(parsed_docx, total_seconds=480)
    roles = {slide.role for _, slide in deck.iter_slides()}
    assert roles & {SlideRole.BIG_FIGURE, SlideRole.TEXT_FIGURE, SlideRole.TABLE}


def test_planner_marks_extracted_assets_as_extracted(parsed_docx):
    """Provenance must record that figures came from the author, not a model."""
    deck = plan_deck(parsed_docx, total_seconds=480)
    assert deck.assets
    for asset in deck.assets:
        assert asset.origin.value == "extracted"


# --------------------------------------------------------------------------- #
# End to end
# --------------------------------------------------------------------------- #


def test_end_to_end_docx_to_pptx(parsed_docx, tmp_path):
    """A manuscript must reach a rendered deck with nothing overflowing."""
    pptx = pytest.importorskip("pptx")
    from rostrum.budget.allocate import allocate
    from rostrum.render.pptx import export_script, render_pptx
    from rostrum.templates.binding import bind, capacity_caps, overflow_rate
    from rostrum.templates.ingest_pptx import ingest_pptx

    template = tmp_path / "t.pptx"
    pptx.Presentation().save(str(template))

    deck = plan_deck(parsed_docx, total_seconds=480, scenario=Scenario.GRANT_DEFENSE)
    contract, _ = ingest_pptx(str(template), template_id="t", license="test")
    binding = bind(deck, contract)
    allocate(deck, apply=True, capacity=capacity_caps(binding))

    out = tmp_path / "out.pptx"
    report = render_pptx(deck, contract, binding, str(out))

    assert out.exists()
    assert report.slides_written > 0
    assert report.missing_assets == []
    assert overflow_rate(deck, binding) == 0.0

    script_path = tmp_path / "s.md"
    export_script(deck, str(script_path))
    assert script_path.exists()
    body = script_path.read_text(encoding="utf-8")
    # The script is the other half of the deliverable: content the slides could
    # not hold must be speakable, with per-slide timings to rehearse against.
    assert "研究背景与问题" in body
    assert "s)" in body or "秒" in body


def test_binding_avoids_layouts_whose_title_sits_low(parsed_docx, tmp_path):
    """A figure slide must not be titled beneath its own image.

    PowerPoint's "Picture with Caption" is the only exact big_figure match in the
    default template, and its title sits at y=0.7. Visual review caught the
    resulting slide with its heading stranded at the bottom of the page.
    """
    pptx = pytest.importorskip("pptx")
    from rostrum.templates.binding import bind
    from rostrum.templates.ingest_pptx import ingest_pptx

    template = tmp_path / "t.pptx"
    pptx.Presentation().save(str(template))
    contract, _ = ingest_pptx(str(template), template_id="t", license="test")

    deck = plan_deck(parsed_docx, total_seconds=480)
    report = bind(deck, contract)

    layouts = {x.layout_id: x for x in contract.layouts}
    for _, slide in deck.iter_slides():
        binding = report.bindings.get(slide.uid)
        if binding is None:
            continue
        layout = layouts[binding.layout_id]
        titles = [s for s in layout.slots if s.kind == "title"]
        if titles:
            assert min(t.box.y for t in titles) <= 0.45


def test_allocation_never_empties_a_slide():
    """A slide must always show something, however weak its content scores.

    Demotion by importance is correct in aggregate but can strip a page bare: a
    section whose single paragraph falls just under the threshold sends
    everything to the script, leaving the audience reading a bare heading. Found
    by an end-to-end run over a LaTeX paper, not by a unit test of either layer.
    """
    from rostrum.budget.allocate import allocate
    from rostrum.ir.nodes import Block, Deck, DeckMeta, DeliveryPlan, Section, Slide

    weak = Block(
        type=BlockType.BULLET,
        content="一句重要性偏低但仍是本页唯一内容的陈述",
        importance=0.1,
        channel=Channel.SLIDE,
    )
    deck = Deck(
        meta=DeckMeta(title="t"),
        delivery=DeliveryPlan(total_seconds=300),
        sections=[
            Section(title="s", slides=[Slide(title="p", blocks=[weak])])
        ],
    )
    plan = allocate(deck, apply=True)

    slide = deck.sections[0].slides[0]
    assert slide.slide_blocks(), "the last block on a page must not be demoted"
    assert any(a.reinstated for a in plan.slides), "the rescue must be reported"


def test_allocation_reports_reinstatement_rather_than_hiding_it():
    """The user must be able to see that a weak line was kept for lack of any
    stronger one, since that is a signal the section needs rewriting."""
    from rostrum.budget.allocate import allocate
    from rostrum.ir.nodes import Block, Deck, DeckMeta, DeliveryPlan, Section, Slide

    deck = Deck(
        meta=DeckMeta(title="t"),
        delivery=DeliveryPlan(total_seconds=300),
        sections=[
            Section(
                title="s",
                slides=[
                    Slide(
                        title="p",
                        blocks=[
                            Block(
                                type=BlockType.BULLET,
                                content="低分内容",
                                importance=0.08,
                            )
                        ],
                    )
                ],
            )
        ],
    )
    plan = allocate(deck, apply=True)
    rescued = [uid for a in plan.slides for uid in a.reinstated]
    assert len(rescued) == 1
