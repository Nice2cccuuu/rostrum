"""Tests for natural-language revision.

Three properties are load-bearing here, and each has its own section below.

**Containment.** An operation may only touch what it declared. This is verified
against a deliberately malicious operation that under-reports its blast radius,
because the assertion is worthless if it has never been shown to fire.

**Honest refusal.** An utterance the interpreter cannot resolve must produce a
question, not a guess. Several tests assert that a patch is *not* produced --
silently applying a plausible misreading is the failure mode that makes this kind
of tool untrustworthy, and it is worse than admitting confusion.

**Provenance survives editing.** Rewriting keeps the source spans and downgrades
the derivation; merging unions the spans; hand-typed text stops claiming to be
the author's words. A deck whose selling point is traceability cannot lose that
property the moment someone edits it.
"""

from __future__ import annotations

from typing import Literal

import pytest

from rostrum.ir.enums import Channel, Density, Derivation, Scenario
from rostrum.patch.apply import ContainmentError, PatchError, apply_patch
from rostrum.patch.interpret import interpret
from rostrum.patch.ops import (
    DeleteBlock,
    EditLog,
    InsertBlock,
    MergeBlocks,
    Patch,
    Pin,
    Retime,
    Rewrite,
    SetChannel,
    SetDwell,
    SetText,
    SplitSlide,
    _Op,
)
from rostrum.patch.session import Session, diff_decks

pytest.importorskip("docx")


# --------------------------------------------------------------------------- #
# Fixtures
# --------------------------------------------------------------------------- #


@pytest.fixture(scope="module")
def manuscript(tmp_path_factory):
    import docx

    d = tmp_path_factory.mktemp("src")
    doc = docx.Document()
    doc.core_properties.title = "低资源表征学习"
    doc.core_properties.author = "张三"

    doc.add_heading("低资源表征学习", 1)
    doc.add_heading("研究背景与问题", 1)
    doc.add_paragraph(
        "深度表征学习在大规模标注数据上取得了显著成功，但在医学影像等真实场景中，"
        "获取高质量标注的成本极高，可用样本常低于千级规模。"
    )
    doc.add_paragraph(
        "现有自监督方法在小样本条件下普遍出现表征坍缩，即不同类别的样本被映射到"
        "相近的表征空间区域，导致下游性能急剧下降。"
    )
    doc.add_paragraph("本项目拟解决的核心科学问题是如何构造不坍缩的表征空间。")
    doc.add_heading("研究目标与内容", 1)
    doc.add_paragraph("总体思路是以结构先验约束表征空间的几何形态。")
    doc.add_paragraph("研究内容一：结构一致性正则项的设计", style="List Bullet")
    doc.add_paragraph("研究内容二：小样本几何约束方法", style="List Bullet")
    doc.add_heading("创新点", 1)
    doc.add_paragraph("提出结构一致性正则并给出理论刻画", style="List Number")
    doc.add_paragraph("参数量降低54%，精度提升7.5个百分点", style="List Number")
    doc.add_paragraph("给出可迁移至三类下游任务的统一框架", style="List Number")
    doc.add_heading("研究基础", 1)
    doc.add_paragraph("前期工作已在三个公开数据集上验证了上述结论。")
    doc.add_heading("可行性分析", 1)
    doc.add_paragraph("团队具备完整的算力条件与数据积累。")

    path = d / "m.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def deck(manuscript, tmp_path):
    from rostrum.budget.allocate import allocate
    from rostrum.ingest.docx_parser import parse_docx
    from rostrum.ingest.planner import plan_deck

    doc = parse_docx(manuscript, asset_dir=str(tmp_path / "a"))
    d = plan_deck(doc, total_seconds=480, scenario=Scenario.GRANT_DEFENSE)
    allocate(d, apply=True)
    return d


def _body_blocks(deck):
    return [
        b
        for _, _, b in deck.iter_blocks()
        if b.channel is Channel.SLIDE and not b.is_visual
    ]


def _slide_titled(deck, fragment):
    for _, slide in deck.iter_slides():
        if fragment in (slide.title or ""):
            return slide
    raise AssertionError(f"no slide titled like {fragment!r}")


# --------------------------------------------------------------------------- #
# Containment: the property that makes this trustworthy
# --------------------------------------------------------------------------- #


class _Sneaky(_Op):
    """An op that edits a node it never declared. Exists only to be caught."""

    op: Literal["sneaky"] = "sneaky"
    victim: str


def test_containment_catches_an_op_that_exceeds_its_blast_radius(deck):
    """The assertion is worthless unless it has been shown to fire."""
    import rostrum.patch.apply as apply_module

    blocks = _body_blocks(deck)

    def handler(d, op, report, cap):
        d.find(op.target).content = "declared"
        d.find(op.victim).content = "undeclared"

    apply_module._HANDLERS["sneaky"] = handler
    try:
        patch = Patch.model_construct(
            patch_id="evil",
            operations=[_Sneaky(target=blocks[0].uid, victim=blocks[1].uid)],
        )
        with pytest.raises(ContainmentError, match="blast radius"):
            apply_patch(deck, patch)
    finally:
        del apply_module._HANDLERS["sneaky"]


def test_a_failed_patch_leaves_the_original_untouched(deck):
    blocks = _body_blocks(deck)
    before = blocks[0].content
    patch = Patch(
        patch_id="bad",
        operations=[
            SetText(target=blocks[0].uid, value="新内容"),
            SetText(target="blk_does_not_exist", value="x"),
        ],
    )
    with pytest.raises(PatchError, match="no such node"):
        apply_patch(deck, patch)
    assert deck.find(blocks[0].uid).content == before


def test_apply_never_mutates_its_input(deck):
    """Undo depends on this: history is kept as values, not inverses."""
    blocks = _body_blocks(deck)
    original = blocks[0].content
    patch = Patch(
        patch_id="p", operations=[SetText(target=blocks[0].uid, value="改过了")]
    )
    after, _ = apply_patch(deck, patch)
    assert after.find(blocks[0].uid).content == "改过了"
    assert deck.find(blocks[0].uid).content == original


def test_containment_allows_the_parent_of_a_removed_block(deck):
    """Deleting a block necessarily changes its slide; that is not a violation."""
    blocks = _body_blocks(deck)
    patch = Patch(
        patch_id="p", operations=[DeleteBlock(target=blocks[0].uid, hard=True)]
    )
    after, report = apply_patch(deck, patch)
    assert blocks[0].uid in report.removed_uids
    assert after.find(blocks[0].uid) is None


# --------------------------------------------------------------------------- #
# Provenance survives editing
# --------------------------------------------------------------------------- #


def test_rewrite_keeps_spans_and_marks_the_text_compressed(deck):
    from rostrum.budget.allocate import count_units

    # Shortening cuts only at a clause boundary, so the target must *have* one.
    # Picking "the longest block" no longer guarantees that: the planner now
    # splits long paragraphs into single-clause points, which is the desired
    # outcome and leaves nothing multi-clause behind. So the test supplies its own
    # target instead of depending on what planning happens to leave.
    # Not a pinned block: pinning exempts content from automatic modification by
    # design, so a pinned target makes shortening a legitimate no-op.
    candidates = [b for b in _body_blocks(deck) if not b.pinned]
    assert candidates, "fixture has no unpinned bullet"
    long_block = candidates[0]
    long_block.content = (
        "深度表征学习在大规模标注数据上取得了显著成功，"
        "但在医学影像等真实场景中获取标注的成本极高。"
    )
    spans_before = len(long_block.spans)
    limit = max(8, count_units(long_block.content) // 2)
    original = long_block.content

    patch = Patch(
        patch_id="p",
        operations=[
            Rewrite(target=long_block.uid, instruction="短一些", max_units=limit)
        ],
    )
    after, _ = apply_patch(deck, patch)
    edited = after.find(long_block.uid)

    assert count_units(edited.content) < count_units(original)
    assert len(edited.spans) == spans_before, "shortening must not drop the source"
    assert edited.derivation is Derivation.COMPRESSED


def test_rewrite_never_cuts_mid_phrase(deck):
    """A bullet ending in half a clause reads as a bug to the audience."""
    from rostrum.budget.allocate import count_units

    long_block = max(_body_blocks(deck), key=lambda b: count_units(b.content))
    patch = Patch(
        patch_id="p",
        operations=[Rewrite(target=long_block.uid, instruction="短", max_units=20)],
    )
    after, _ = apply_patch(deck, patch)
    text = after.find(long_block.uid).content
    assert not text.endswith(("，", ",", "、", "；"))


def test_rewrite_reports_when_it_cannot_meet_the_budget(deck):
    """Returning the original silently would look like the edit was applied."""
    long_block = max(_body_blocks(deck), key=lambda b: len(b.content))
    patch = Patch(
        patch_id="p",
        operations=[Rewrite(target=long_block.uid, instruction="短", max_units=3)],
    )
    after, report = apply_patch(deck, patch)
    assert after.find(long_block.uid).content == long_block.content
    assert any("could not be shortened" in n for n in report.notes)


def test_hand_typed_text_stops_claiming_to_be_the_authors_words(deck):
    verbatim = next(
        b for b in _body_blocks(deck) if b.derivation is Derivation.VERBATIM
    )
    patch = Patch(
        patch_id="p", operations=[SetText(target=verbatim.uid, value="我自己写的一句话")]
    )
    after, _ = apply_patch(deck, patch)
    assert after.find(verbatim.uid).derivation is Derivation.AUTHORED


def test_merging_unions_the_sources_and_stops_claiming_verbatim(deck):
    slide = next(s for _, s in deck.iter_slides() if len(s.blocks) >= 2)
    first, second = slide.blocks[0], slide.blocks[1]
    expected = len({(s.start, s.end) for s in first.spans + second.spans})

    patch = Patch(
        patch_id="p",
        operations=[MergeBlocks(target=first.uid, others=[second.uid])],
    )
    after, report = apply_patch(deck, patch)
    merged = after.find(first.uid)

    assert second.uid in report.removed_uids
    assert after.find(second.uid) is None
    assert len(merged.spans) == expected
    if len(merged.spans) > 1:
        assert merged.derivation is Derivation.SYNTHESIZED


def test_inserted_text_may_not_pose_as_sourced_content(deck):
    slide = next(s for _, s in deck.iter_slides() if s.blocks)
    patch = Patch(
        patch_id="p",
        operations=[
            InsertBlock(
                target=slide.uid,
                block={
                    "type": "bullet",
                    "content": "凭空加的一条",
                    "derivation": "verbatim",
                    "spans": [],
                },
            )
        ],
    )
    with pytest.raises(PatchError, match="authored"):
        apply_patch(deck, patch)


# --------------------------------------------------------------------------- #
# Routing: content leaves the slide but not the deck
# --------------------------------------------------------------------------- #


def test_moving_content_to_the_script_keeps_the_text(deck):
    """The core promise: what comes off the slide is spoken, not lost."""
    block = _body_blocks(deck)[0]
    patch = Patch(
        patch_id="p",
        operations=[SetChannel(target=block.uid, channel=Channel.SCRIPT)],
    )
    after, _ = apply_patch(deck, patch)
    moved = after.find(block.uid)
    assert moved is not None
    assert moved.channel is Channel.SCRIPT
    assert moved.content == block.content


def test_a_figure_cannot_be_moved_to_the_spoken_script(deck):
    visual = next(
        (b for _, _, b in deck.iter_blocks() if b.is_visual), None
    )
    if visual is None:
        pytest.skip("manuscript has no figure or table")
    patch = Patch(
        patch_id="p",
        operations=[SetChannel(target=visual.uid, channel=Channel.SCRIPT)],
    )
    with pytest.raises(PatchError, match="cannot be moved"):
        apply_patch(deck, patch)


def test_a_soft_delete_is_reversible(deck):
    block = _body_blocks(deck)[0]
    patch = Patch(patch_id="p", operations=[DeleteBlock(target=block.uid)])
    after, _ = apply_patch(deck, patch)
    assert after.find(block.uid).channel is Channel.DROP


# --------------------------------------------------------------------------- #
# Timing
# --------------------------------------------------------------------------- #


def test_a_stated_dwell_survives_the_next_reallocation(deck):
    """Otherwise "give this page 30 more seconds" is undone by the next edit.

    The allocator used to overwrite dwell unconditionally, so the patch appeared
    to work and the diff came back empty.
    """
    slide = max(
        (s for _, s in deck.iter_slides() if s.blocks),
        key=lambda s: s.dwell_seconds or 0,
    )
    fixed = Patch(patch_id="p1", operations=[SetDwell(target=slide.uid, seconds=90)])
    after, _ = apply_patch(deck, fixed)
    assert after.find(slide.uid).dwell_seconds == pytest.approx(90, abs=1)
    assert after.find(slide.uid).dwell_locked

    shrink = Patch(
        patch_id="p2", operations=[Retime(target=after.uid, total_seconds=300)]
    )
    after2, _ = apply_patch(after, shrink)
    assert after2.find(slide.uid).dwell_seconds == pytest.approx(90, abs=1)


def test_retime_redistributes_the_remaining_time(deck):
    patch = Patch(
        patch_id="p", operations=[Retime(target=deck.uid, total_seconds=300)]
    )
    after, _ = apply_patch(deck, patch)
    assert after.delivery.total_seconds == 300
    total = sum(s.dwell_seconds or 0 for _, s in after.iter_slides())
    assert total <= 300 * 1.05


def test_retime_must_target_the_deck(deck):
    slide = next(s for _, s in deck.iter_slides())
    patch = Patch(
        patch_id="p", operations=[Retime(target=slide.uid, total_seconds=300)]
    )
    with pytest.raises(PatchError, match="targets the deck"):
        apply_patch(deck, patch)


def test_pinned_content_survives_re_budgeting(deck):
    """The contract that makes iterative editing workable."""
    block = _body_blocks(deck)[0]
    after, _ = apply_patch(
        deck, Patch(patch_id="p1", operations=[Pin(target=block.uid, pinned=True)])
    )
    after2, _ = apply_patch(
        after,
        Patch(patch_id="p2", operations=[Retime(target=after.uid, total_seconds=240)]),
    )
    survivor = after2.find(block.uid)
    assert survivor.pinned
    assert survivor.channel is Channel.SLIDE


def test_the_allocator_flags_a_slide_that_eats_the_talk(deck):
    """Pacing advice the user can act on with one sentence."""
    from rostrum.budget.allocate import allocate

    plan = allocate(deck, apply=True)
    shares = [
        a.dwell_seconds / max(1.0, sum(x.dwell_seconds for x in plan.slides))
        for a in plan.slides
    ]
    if max(shares) < 0.22:
        pytest.skip("no slide dominates this manuscript")
    assert any("of the talk" in n for n in plan.notes)


# --------------------------------------------------------------------------- #
# Structure
# --------------------------------------------------------------------------- #


def test_splitting_a_slide_keeps_every_block(deck):
    slide = next(s for _, s in deck.iter_slides() if len(s.blocks) >= 2)
    before = [b.uid for b in slide.blocks]
    patch = Patch(
        patch_id="p",
        operations=[SplitSlide(target=slide.uid, after_block=slide.blocks[0].uid)],
    )
    after, report = apply_patch(deck, patch)
    assert len(report.created_uids) == 1
    surviving = {b.uid for _, _, b in after.iter_blocks()}
    assert set(before) <= surviving


def test_splitting_at_the_last_block_is_refused(deck):
    slide = next(s for _, s in deck.iter_slides() if len(s.blocks) >= 2)
    patch = Patch(
        patch_id="p",
        operations=[SplitSlide(target=slide.uid, after_block=slide.blocks[-1].uid)],
    )
    with pytest.raises(PatchError, match="nothing to move"):
        apply_patch(deck, patch)


def test_a_block_operation_refuses_a_slide_target(deck):
    slide = next(s for _, s in deck.iter_slides())
    patch = Patch(patch_id="p", operations=[SetText(target=slide.uid, value="x")])
    with pytest.raises(PatchError, match="targets a block"):
        apply_patch(deck, patch)


# --------------------------------------------------------------------------- #
# Interpretation: what it understands
# --------------------------------------------------------------------------- #


def test_understands_moving_the_last_n_bullets_to_the_script(deck):
    slide = _slide_titled(deck, "研究背景")
    index = [s for _, s in deck.iter_slides()].index(slide) + 1

    result = interpret(f"第{index}页太满了，把后两条放到讲稿", deck)
    assert result.ok
    assert result.confidence >= 0.75
    ops = result.patch.operations
    assert len(ops) == 2, "「后两条」 states a count; it must be honoured"
    assert all(o.op == "set_channel" for o in ops)


def test_spoken_chinese_uses_liang_not_er_for_quantities(deck):
    """"两条" is how people say it; "二条" is how nobody says it."""
    from rostrum.patch.interpret import _count_phrase

    assert _count_phrase("把后两条放到讲稿") == 2
    assert _count_phrase("最后三个点") == 3


def test_understands_a_deck_wide_duration_change(deck):
    result = interpret("整体压到8分钟", deck)
    assert result.ok
    assert result.confidence >= 0.9
    op = result.patch.operations[0]
    assert op.op == "retime"
    assert op.total_seconds == 480


def test_understands_giving_one_page_more_time(deck):
    result = interpret("创新点这页多给30秒", deck)
    assert result.ok
    op = result.patch.operations[0]
    assert op.op == "set_dwell"
    target = deck.find(op.target)
    assert "创新点" in target.title


def test_a_literal_substitution_also_covers_titles(deck):
    """Rewriting body text but not the heading put both wordings on one page.

    The phrase is taken from a title that actually exists in this deck, so the
    test proves the behaviour rather than depending on the fixture's wording.
    """
    titled = next(s for _, s in deck.iter_slides() if s.title and len(s.title) > 3)
    word = titled.title[:3]

    result = interpret(f"把「{word}」改成「甲乙丙」", deck)
    assert result.ok
    kinds = {o.op for o in result.patch.operations}
    assert "set_title" in kinds, "a heading is user-visible text too"


def test_understands_a_density_preference(deck):
    result = interpret("整体做稀疏一点", deck)
    assert result.ok
    op = result.patch.operations[0]
    assert op.op == "retime"
    assert op.density is Density.SPARSE


def test_a_selection_resolves_this_without_any_description(deck):
    """Click-to-select and language must feed the same mechanism."""
    block = _body_blocks(deck)[0]
    result = interpret("把这条改短一点", deck, selection=[block.uid])
    assert result.ok
    assert result.confidence >= 0.85
    assert result.patch.operations[0].target == block.uid


def test_the_users_words_are_kept_on_every_operation(deck):
    """An edit history is useless if it cannot show why something changed."""
    result = interpret("整体压到8分钟", deck)
    assert result.patch.utterance == "整体压到8分钟"
    assert all(o.rationale for o in result.patch.operations)


# --------------------------------------------------------------------------- #
# Interpretation: what it refuses
# --------------------------------------------------------------------------- #


def test_an_unresolvable_request_asks_instead_of_guessing(deck):
    result = interpret("把这页弄好看点", deck)
    assert not result.ok
    assert result.question, "confusion must be voiced, not resolved by guessing"


def test_a_page_number_out_of_range_is_reported_not_clamped(deck):
    result = interpret("第99页太满了", deck)
    assert not result.ok
    assert "99" in result.reason or "页" in result.reason


def test_a_missing_phrase_is_reported_rather_than_approximated(deck):
    result = interpret("把「量子计算」改成「经典计算」", deck)
    assert not result.ok
    assert "量子计算" in result.reason


def test_a_retitle_without_a_new_title_asks_for_one(deck):
    result = interpret("第3页标题改一下", deck)
    assert not result.ok
    assert result.question


def test_deleting_without_saying_what_asks_which(deck):
    slide = _slide_titled(deck, "创新点")
    index = [s for _, s in deck.iter_slides()].index(slide) + 1
    result = interpret(f"第{index}页删掉一条", deck)
    assert not result.ok
    assert result.question


def test_an_empty_utterance_is_not_an_error(deck):
    result = interpret("   ", deck)
    assert not result.ok
    assert not result.patch


# --------------------------------------------------------------------------- #
# Sessions: undo, redo, replay
# --------------------------------------------------------------------------- #


def test_undo_restores_the_previous_state(deck):
    session = Session(original=deck)
    before = session.current.delivery.total_seconds

    session.say("整体压到6分钟")
    assert session.current.delivery.total_seconds == 360

    session.undo()
    assert session.current.delivery.total_seconds == before


def test_redo_reapplies_what_undo_removed(deck):
    session = Session(original=deck)
    session.say("整体压到6分钟")
    session.undo()
    session.redo()
    assert session.current.delivery.total_seconds == 360


def test_undo_on_an_untouched_session_is_a_no_op(deck):
    session = Session(original=deck)
    assert session.undo() is None


def test_a_new_edit_clears_the_redo_stack(deck):
    session = Session(original=deck)
    session.say("整体压到6分钟")
    session.undo()
    session.say("整体压到5分钟")
    assert session.redo() is None


def test_replaying_a_log_reproduces_the_deck(deck):
    """Reproducibility is what makes an edit history worth storing."""
    first = Session(original=deck)
    for utterance in ("整体压到6分钟", "整体做稀疏一点"):
        first.say(utterance)

    stored = EditLog.model_validate_json(first.log.model_dump_json())
    second = Session(original=deck)
    second.replay(stored)

    assert second.current.model_dump_json() == first.current.model_dump_json()


def test_a_log_from_another_deck_is_refused(deck):
    session = Session(original=deck)
    foreign = EditLog(deck_uid="deck_somewhere_else")
    with pytest.raises(PatchError, match="belongs to deck"):
        session.replay(foreign)


def test_history_shows_what_the_user_actually_said(deck):
    session = Session(original=deck)
    session.say("整体压到6分钟")
    assert any("整体压到6分钟" in line for line in session.history())


def test_a_low_confidence_edit_is_previewed_rather_than_applied(deck):
    """The asymmetry that makes the tool both usable and trustworthy."""
    session = Session(original=deck)
    _result, _diff, report = session.say("把这页弄好看点", threshold=0.75)
    assert report is None
    assert not session.log.patches


# --------------------------------------------------------------------------- #
# Diffs
# --------------------------------------------------------------------------- #


def test_a_diff_shows_the_change_the_user_asked_for(deck):
    """Reallocation ripple must not bury the one line that was requested.

    Asking for 30 more seconds on a page produced an empty-looking diff, because
    the requested change was folded into "9 slides were re-timed".
    """
    session = Session(original=deck)
    result, diff = session.preview("创新点这页多给30秒")
    assert result.ok
    assert diff, "the requested change must be visible"
    assert any(c.kind == "timing" for c in diff.changes)


def test_incidental_timing_ripple_is_summarised(deck):
    session = Session(original=deck)
    _, diff = session.preview("创新点这页多给30秒")
    assert diff.notes, "collateral re-timing should be summarised, not listed"


def test_a_channel_change_reads_as_a_move_not_a_deletion(deck):
    """Wording matters: the user must see the content was kept."""
    block = _body_blocks(deck)[0]
    patch = Patch(
        patch_id="p",
        operations=[SetChannel(target=block.uid, channel=Channel.SCRIPT)],
    )
    after, _ = apply_patch(deck, patch)
    diff = diff_decks(deck, after)
    change = next(c for c in diff.changes if c.kind == "channel")
    assert "演讲文稿" in (change.after or "")


def test_preview_does_not_change_the_session(deck):
    session = Session(original=deck)
    before = session.current.model_dump_json()
    session.preview("整体压到6分钟")
    assert session.current.model_dump_json() == before
    assert not session.log.patches


# --------------------------------------------------------------------------- #
# End to end
# --------------------------------------------------------------------------- #


def test_a_full_revision_round_trip_renders(deck, tmp_path):
    """Say four things, then produce slides and a script that still hold."""
    from rostrum.render import render_pptx
    from rostrum.templates import bind, capacity_caps, ingest_pptx, overflow_rate
    from rostrum.themes import DEFAULT_THEME_ID, build_template, get_theme

    template = tmp_path / "t.pptx"
    build_template(get_theme(DEFAULT_THEME_ID), str(template))
    contract, _ = ingest_pptx(str(template), template_id="t", license="builtin")

    session = Session(
        original=deck, capacity=capacity_caps(bind(deck, contract))
    )
    for utterance in (
        "创新点这页拆成两页",
        "把「本项目」改成「本课题」",
        "整体压到6分钟",
    ):
        session.say(utterance)

    revised = session.current
    binding = bind(revised, contract)
    out = tmp_path / "out.pptx"
    report = render_pptx(revised, contract, binding, str(out))

    assert out.exists()
    assert report.slides_written >= 8
    assert overflow_rate(revised, binding) == 0.0
    assert revised.delivery.total_seconds == 360
    assert "本项目" not in " ".join(
        (s.title or "") + " ".join(b.content for b in s.blocks)
        for _, s in revised.iter_slides()
    )


def test_content_moved_off_the_slide_appears_in_the_script(deck, tmp_path):
    """The promise that makes decluttering safe."""
    from rostrum.render import export_script

    slide = _slide_titled(deck, "研究背景")
    index = [s for _, s in deck.iter_slides()].index(slide) + 1

    session = Session(original=deck)
    result, _, report = session.say(f"第{index}页太满了，把后两条放到讲稿")
    assert report is not None

    moved = [
        deck.find(op.target).content
        for op in result.patch.operations
        if op.op == "set_channel"
    ]
    script = export_script(session.current, str(tmp_path / "s.md"))
    for text in moved:
        assert text[:20] in script, "content left the slide but not the talk"
