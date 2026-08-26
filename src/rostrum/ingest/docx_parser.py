"""DOCX parser.

Reads a Word manuscript into a :class:`ParsedDocument`: heading hierarchy,
paragraphs, lists, tables, and the author's own embedded images extracted as
original bytes.

Extraction detail worth stating: images come out of the OPC package byte for
byte, not re-encoded and not screenshotted. A figure in a grant proposal is
evidence; re-compressing it degrades a reviewer's ability to read axis labels,
and regenerating it would be indefensible.
"""

from __future__ import annotations

import contextlib
import os
import re

from rostrum.ingest.model import (
    ExtractedAsset,
    ParsedDocument,
    Segment,
    SegmentKind,
    TextBuilder,
)
from rostrum.ir.enums import AssetKind

try:
    import docx
    from docx.table import Table

    _HAVE_DOCX = True
except ImportError:  # pragma: no cover - optional dependency
    _HAVE_DOCX = False


# Caption paragraphs, in both conventions academics actually use.
_CAPTION_RE = re.compile(
    r"^\s*(figure|fig\.?|table|tab\.?|chart|scheme|图|表|附图|附表)\s*"
    r"([0-9]+|[一二三四五六七八九十]+)?\s*[.:：、]?\s*",
    re.IGNORECASE,
)
_REFERENCE_HEADING = re.compile(
    r"^\s*(references?|bibliography|参考文献|引用文献)\s*$", re.IGNORECASE
)
_EQUATION_HINT = re.compile(r"[=≈≤≥∑∫∂√±×÷∞]|\\[a-zA-Z]{2,}")

# Style names that mark a caption regardless of the paragraph's text.
_CAPTION_STYLES = {"caption", "图注", "表注", "图表标题"}
_QUOTE_STYLES = {"quote", "intense quote", "block text", "引用"}
_CODE_STYLES = {"code", "source code", "html preformatted", "代码"}


def parse_docx(
    path: str,
    *,
    doc_id: str = "manuscript",
    language: str = "zh",
    asset_dir: str | None = None,
) -> ParsedDocument:
    """Parse ``path`` into a :class:`ParsedDocument`.

    Parameters
    ----------
    asset_dir:
        Directory for extracted images. Defaults to ``<path>_assets``. Figures
        are written here with their original bytes preserved.
    """
    if not _HAVE_DOCX:  # pragma: no cover
        raise RuntimeError("python-docx is required: pip install python-docx")

    document = docx.Document(path)
    builder = TextBuilder()
    assets: list[ExtractedAsset] = []
    warnings: list[str] = []

    asset_dir = asset_dir or f"{os.path.splitext(path)[0]}_assets"
    images = _extract_images(document, asset_dir, warnings)

    section_stack: list[str] = []
    in_references = False
    pending_caption: Segment | None = None
    image_cursor = 0
    table_index = 0

    for block in _iter_blocks(document):
        if isinstance(block, Table):
            table_index += 1
            asset = _table_asset(block, table_index)
            seg = builder.add(
                _table_preview(block),
                SegmentKind.TABLE,
                section_path=tuple(section_stack),
                asset_id=asset.asset_id,
            )
            if seg is not None:
                asset.spans.append(seg.span(doc_id))
                if pending_caption is not None:
                    asset.caption = pending_caption.text
                    pending_caption.asset_id = asset.asset_id
                    pending_caption = None
            assets.append(asset)
            continue

        text = block.text.strip()
        style = (block.style.name if block.style is not None else "") or ""
        style_key = style.strip().lower()

        # A paragraph holding an inline image contributes the image, not text.
        n_images = _count_inline_images(block)
        if n_images:
            for _ in range(n_images):
                if image_cursor >= len(images):
                    break
                asset = images[image_cursor]
                image_cursor += 1
                anchor = builder.add(
                    text or f"[{asset.source_label or asset.asset_id}]",
                    SegmentKind.CAPTION if text else SegmentKind.PARAGRAPH,
                    section_path=tuple(section_stack),
                    asset_id=asset.asset_id,
                )
                if anchor is not None:
                    asset.spans.append(anchor.span(doc_id))
                if text:
                    asset.caption = text
                assets.append(asset)
            if text and _is_caption(text, style_key):
                continue
            if not text:
                continue

        if not text:
            continue

        heading_level = _heading_level(block, style_key)
        if heading_level:
            in_references = bool(_REFERENCE_HEADING.match(text))
            del section_stack[heading_level - 1 :]
            section_stack.append(text)
            builder.add(
                text,
                SegmentKind.HEADING,
                level=heading_level,
                section_path=tuple(section_stack[:-1]),
                style=style,
            )
            pending_caption = None
            continue

        if in_references:
            builder.add(
                text,
                SegmentKind.REFERENCE,
                section_path=tuple(section_stack),
                style=style,
            )
            continue

        kind, level = _classify(block, text, style_key)
        seg = builder.add(
            text,
            kind,
            level=level,
            section_path=tuple(section_stack),
            style=style,
        )

        if kind is SegmentKind.CAPTION and seg is not None:
            # Academic convention puts figure captions *below* the figure and
            # table captions *above* the table, so both directions must bind.
            # A trailing caption claims the most recent unlabelled asset.
            trailing = _claim_trailing(assets, seg, doc_id)
            pending_caption = None if trailing else seg
        else:
            pending_caption = None

    if image_cursor < len(images):
        # Images that no paragraph anchored: keep them, but say so. Losing an
        # author's figure silently would be worse than an unplaced one.
        for asset in images[image_cursor:]:
            assets.append(asset)
        warnings.append(
            f"{len(images) - image_cursor} image(s) had no anchoring paragraph "
            "and carry no source span"
        )

    parsed = ParsedDocument(
        doc_id=doc_id,
        text=builder.text,
        segments=builder.segments,
        assets=assets,
        title=_document_title(document, builder.segments),
        authors=_document_authors(document),
        source_path=path,
        source_format="docx",
        language=language,
        warnings=warnings,
    )
    parsed.warnings.extend(parsed.verify())
    return parsed


def _claim_trailing(
    assets: list[ExtractedAsset], caption: Segment, doc_id: str
) -> bool:
    """Bind a caption to the asset it follows, if that asset still lacks one.

    Only the most recent asset is considered, and only when it is uncaptioned:
    a caption further away is more likely prose that merely mentions a figure.
    """
    for asset in reversed(assets):
        if asset.caption:
            return False
        # Match the label kind so a figure caption cannot claim a table.
        wants_table = bool(re.match(r"^\s*(table|tab\.?|表)", caption.text, re.I))
        is_table = asset.kind is AssetKind.TABLE
        if wants_table != is_table:
            return False
        asset.caption = caption.text
        caption.asset_id = asset.asset_id
        asset.spans.append(caption.span(doc_id))
        return True
    return False


# --------------------------------------------------------------------------- #
# Block iteration
# --------------------------------------------------------------------------- #


def _iter_blocks(document):
    """Yield paragraphs and tables in true document order.

    ``document.paragraphs`` and ``document.tables`` are separate sequences, so
    iterating them in turn scrambles the order and destroys the relationship
    between a caption and the table it introduces. Walking the body XML is the
    only way to preserve it.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _count_inline_images(paragraph) -> int:
    from docx.oxml.ns import qn

    return len(paragraph._p.findall(f".//{qn('a:blip')}"))


# --------------------------------------------------------------------------- #
# Classification
# --------------------------------------------------------------------------- #


def _heading_level(paragraph, style_key: str) -> int:
    """Heading depth, or 0 for body text.

    Style name first, then outline level: a template may define its own heading
    styles that do not begin with "Heading".
    """
    match = re.match(r"^heading (\d)$", style_key)
    if match:
        return int(match.group(1))
    if style_key in {"title", "标题"}:
        return 1
    match = re.match(r"^标题\s*(\d)$", style_key)
    if match:
        return int(match.group(1))

    with contextlib.suppress(AttributeError, KeyError, ValueError):
        fmt = paragraph._p.pPr
        if fmt is not None and fmt.outlineLvl is not None:
            return int(fmt.outlineLvl.val) + 1
    return 0


def _classify(paragraph, text: str, style_key: str) -> tuple[SegmentKind, int]:
    if _is_caption(text, style_key):
        return SegmentKind.CAPTION, 0
    if style_key in _QUOTE_STYLES:
        return SegmentKind.QUOTE, 0
    if style_key in _CODE_STYLES:
        return SegmentKind.CODE, 0

    level = _list_level(paragraph, style_key)
    if level is not None:
        return SegmentKind.LIST_ITEM, level

    # A short line dense in mathematical operators is a display equation.
    if len(text) < 200 and _EQUATION_HINT.search(text):
        letters = sum(1 for c in text if c.isalpha() or "\u4e00" <= c <= "\u9fff")
        if letters < len(text) * 0.55:
            return SegmentKind.EQUATION, 0

    return SegmentKind.PARAGRAPH, 0


def _is_caption(text: str, style_key: str) -> bool:
    if style_key in _CAPTION_STYLES:
        return True
    # Require a label-like opening *and* brevity: a paragraph beginning "Table 1
    # shows that..." is prose about a table, not its caption.
    return bool(_CAPTION_RE.match(text)) and len(text) < 200


def _list_level(paragraph, style_key: str) -> int | None:
    if "list" in style_key or "列表" in style_key:
        with contextlib.suppress(AttributeError, KeyError, ValueError):
            numpr = paragraph._p.pPr.numPr
            if numpr is not None and numpr.ilvl is not None:
                return int(numpr.ilvl.val)
        return 0
    with contextlib.suppress(AttributeError, KeyError, ValueError):
        numpr = paragraph._p.pPr.numPr
        if numpr is not None:
            return int(numpr.ilvl.val) if numpr.ilvl is not None else 0
    # Manual bullets, which are extremely common in real manuscripts.
    if re.match(r"^\s*[-•·▪◦*]\s+", paragraph.text):
        return 0
    if re.match(r"^\s*(\(\d+\)|\d+[.)、]|[（(][一二三四五六七八九十][)）])\s+", paragraph.text):
        return 0
    return None


# --------------------------------------------------------------------------- #
# Assets
# --------------------------------------------------------------------------- #


def _extract_images(document, asset_dir: str, warnings: list[str]):
    """Write embedded images out with original bytes preserved."""
    assets: list[ExtractedAsset] = []
    parts = [
        part
        for rel_id, part in sorted(document.part.related_parts.items())
        if getattr(part, "content_type", "").startswith("image/")
    ]
    if not parts:
        return assets

    os.makedirs(asset_dir, exist_ok=True)
    for i, part in enumerate(parts, 1):
        ext = os.path.splitext(part.partname)[1] or ".png"
        out = os.path.join(asset_dir, f"figure-{i:02d}{ext}")
        try:
            with open(out, "wb") as fh:
                fh.write(part.blob)  # byte-for-byte, never re-encoded
        except OSError as exc:  # pragma: no cover
            warnings.append(f"could not write {out}: {exc}")
            continue
        assets.append(
            ExtractedAsset(
                kind=AssetKind.FIGURE,
                asset_id=f"fig-{i:02d}",
                path=out,
                source_label=f"Figure {i}",
                intrinsic_aspect=_aspect(out),
            )
        )
    return assets


def _aspect(path: str) -> float | None:
    try:
        from PIL import Image

        with Image.open(path) as im:
            w, h = im.size
        return (w / h) if h else None
    except Exception:  # pragma: no cover - non-raster format
        return None


def _table_asset(table, index: int) -> ExtractedAsset:
    """Capture a table as structured data so it can be re-rendered natively.

    Keeping the values rather than an image is what lets the renderer emit a real
    PowerPoint table that a co-author can edit.
    """
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    columns: list[str] = []
    if rows and _looks_like_header(rows[0], rows[1:]):
        columns, rows = rows[0], rows[1:]
    return ExtractedAsset(
        kind=AssetKind.TABLE,
        asset_id=f"tbl-{index:02d}",
        data={"columns": columns, "rows": rows},
        source_label=f"Table {index}",
    )


def _looks_like_header(first: list[str], rest: list[list[str]]) -> bool:
    """Heuristic: a header row is non-numeric while the body is not."""
    if not rest:
        return False
    if any(not c.strip() for c in first):
        return False
    numeric_in_first = sum(1 for c in first if _numeric(c))
    numeric_below = sum(1 for row in rest for c in row if _numeric(c))
    return numeric_in_first == 0 and numeric_below > 0


def _numeric(cell: str) -> bool:
    return bool(re.fullmatch(r"[-+]?[\d.,]+%?", cell.strip()))


def _table_preview(table) -> str:
    """A short textual stand-in so the table occupies a span in the text."""
    rows = [[c.text.strip() for c in row.cells] for row in table.rows]
    if not rows:
        return ""
    head = " | ".join(rows[0])
    return f"[表格 {len(rows)}×{len(rows[0])}] {head}"[:300]


# --------------------------------------------------------------------------- #
# Metadata
# --------------------------------------------------------------------------- #


def _document_title(document, segments: list[Segment]) -> str | None:
    with contextlib.suppress(AttributeError, KeyError):
        title = (document.core_properties.title or "").strip()
        if title:
            return title
    for seg in segments:
        if seg.kind is SegmentKind.HEADING:
            return seg.text
    return None


def _document_authors(document) -> list[str]:
    with contextlib.suppress(AttributeError, KeyError):
        author = (document.core_properties.author or "").strip()
        if author:
            return [a.strip() for a in re.split(r"[;,、；]", author) if a.strip()]
    return []
